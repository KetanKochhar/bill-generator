import os 
import time as t 
import random as r 
import docx as d 
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH as a
document = d.Document()
layout = document.sections[0]
layout.page_width = Mm(80)
layout.page_height = Mm(210)
layout.left_margin = Mm(10)
layout.right_margin = Mm(10)
layout.top_margin = Mm(10)
layout.bottom_margin = Mm(10)
style = document.styles['Normal']
style.paragraph_format.space_after = Mm(1)
def io():
    img = document.add_paragraph()
    document.add_picture("images\\ind.png" , width = Mm(11))
    document.paragraphs[-1].paragraph_format.alignment =  a.CENTER
def bp():
    img = document.add_paragraph()
    document.add_picture("images\\bp.png" , width = Mm(11))
    document.paragraphs[-1].paragraph_format.alignment =  a.CENTER
def hp():
    img = document.add_paragraph()
    document.add_picture("images\\hp.png" , width = Mm(11))
    document.paragraphs[-1].paragraph_format.alignment =  a.CENTER
def rel():
    img = document.add_paragraph()
    document.add_picture("images\\rp.png" , width = Mm(11))
    document.paragraphs[-1].paragraph_format.alignment =  a.CENTER
def makebill():
    print("1. Indian Oil")
    print("2. Bharat petroleum")
    print("3. Hindustan petroleum")
    print("4. Relience petroleum")
    img = int(input("Enter the image number : "))
    name = input("Enter the name of petrol pump : ")
    address = input("Enter the address of petrol pump : ")
    vno = input("Enter the vehicle number : ")
    date = input("Enter he date : ")
    time = input("Enter the time : ")
    fuel = input("Enter the fuel type : ")
    rate = float(input("Enter the rate of fuel : "))
    amt = int(input("Enter the amount of fuel : "))
    billno = r.randint(0,99999)
    tnxID = r.randint(0,999999999999)
    vol = amt/rate
    if img == 1:
        io()
    elif img == 2:
        bp()
    elif img == 3:
        hp()
    elif img == 4:
        rel()
    para1 = document.add_paragraph(name)
    para2 = document.add_paragraph(address)
    para3 = document.add_paragraph("Bill No. : " + str(billno))
    para4 = document.add_paragraph("Trns.ID : " + str(tnxID))
    para5 = document.add_paragraph("Vechicle No. : " + vno)
    para6 = document.add_paragraph("Date : " + date)
    para7 = document.add_paragraph("Time : " + time)
    para8 = document.add_paragraph("Fuel : " + fuel)
    para9 = document.add_paragraph("Rate : " + "INR " + str(rate))
    para10 = document.add_paragraph("Sale : " + "INR " + str(amt))
    para11 = document.add_paragraph("Volume : " + str(vol) +"L" )
    Filename = input("Name of the file : ")
    Fname = Filename + ".docx"
    document.save(Fname)
    print("created succsfully")
    os.system(Fname)
def calculate():
    Break = int(input("Enter the number of breaks : "))
    collect = [ ]
    for x in range(Break):
        first = float(input("Started at : "))
        second = float(input("Ended at : "))
        if first >  second :
            diff = first - second
            print(diff)
            collect.append(diff)
        else :
            diff = second - first 
            print(diff)
            collect.append(diff)
        os.system('cls')
    print("Sum Of data you entered is : ")
    for x in collect:
        sum = 0.0
        a = x
        print(a)
        b = x+1
        print(b)
        sum =+ a+b
    print(sum)
os.system("title Fuel calculator")
os.system("cls")
os.system("color")
os.system("mode con cols=150 lines=30")
for x in range(5):
    if x == 0 :
        print("Loading please wait")
        t.sleep(2)
        os.system("cls")
    if x == 1 :
        print("Loading please wait.")
        t.sleep(2)
        os.system("cls")
    if x == 2 :
        print("Loading please wait..")
        t.sleep(2)
        os.system("cls")
    if x == 3 :
        print("Loading please wait.")
        t.sleep(2)
        os.system("cls")
    if x == 4 :
        print("Loading please wait..")
        t.sleep(2)
        os.system("cls")
    if x == 5 :
        print("Loading please wait...")
        t.sleep(2)
        os.system("cls")
while 1 :
    os.system("cls")
    print("Welcome ")
    print("1. Make a bill ")
    print("2. Calculate the difference")
    print("3. Exit")
    choice = input("Enter your choice : ")
    if choice == "1":
        makebill()
        input("Press Enter to Continue  ")
    elif choice == "2":
        print("!!This is still under process!!")
        calculate()
        input("Press Enter to Continue  ")
    elif choice == "3":
        os.system("cls")
        print("Thank you !")
        t.sleep(2)
        print("closing the program")
        t.sleep(5)
        break
    
    else :
        print("Try Again wrong input")
        input("Press Enter to Continue  ")
exit()
